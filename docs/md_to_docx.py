#!/usr/bin/env python3
"""Convierte los documentos Markdown de HemoPocket a .docx maquetados.
Soporta: títulos (#..####), negrita/cursiva/código, tablas (con estilo de
'bloque de destinatario' sin bordes cuando la cabecera está vacía), listas,
citas, reglas, saltos de página ([SALTO DE PÁGINA]), texto justificado, y
pie de página con numeración."""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BRAND = RGBColor(0xC4, 0x1E, 0x3A)
GREY  = RGBColor(0x66, 0x66, 0x66)
LGREY = RGBColor(0xBB, 0xBB, 0xBB)

INLINE = re.compile(r'(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`|\[[^\]]+?\]\([^)]+?\))')

def add_runs(par, text):
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = par.add_run(part[2:-2]); r.bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = par.add_run(part[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(10)
        elif part.startswith('[') and '](' in part:
            r = par.add_run(part[1:part.index(']')]); r.italic = True
        elif part.startswith('*') and part.endswith('*'):
            r = par.add_run(part[1:-1]); r.italic = True
        else:
            par.add_run(part)

def set_cell_no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e = OxmlElement('w:'+edge); e.set(qn('w:val'),'none'); borders.append(e)
    tblPr.append(borders)

def add_page_number_footer(doc, label):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label + '   ·   Página ')
    run.font.size = Pt(8); run.font.color.rgb = GREY
    # campo PAGE
    fld1 = OxmlElement('w:fldSimple'); fld1.set(qn('w:instr'), 'PAGE')
    r = OxmlElement('w:r'); rpr = OxmlElement('w:rPr')
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'),'16'); rpr.append(sz); r.append(rpr)
    fld1.append(r); p._p.append(fld1)
    run2 = p.add_run(' de '); run2.font.size = Pt(8); run2.font.color.rgb = GREY
    fld2 = OxmlElement('w:fldSimple'); fld2.set(qn('w:instr'), 'NUMPAGES')
    r2 = OxmlElement('w:r'); rpr2 = OxmlElement('w:rPr')
    sz2 = OxmlElement('w:sz'); sz2.set(qn('w:val'),'16'); rpr2.append(sz2); r2.append(rpr2)
    fld2.append(r2); p._p.append(fld2)

def convert(md_path, docx_path, footer_label='HemoPocket'):
    with open(md_path, encoding='utf-8') as f:
        lines = f.read().split('\n')

    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'; normal.font.size = Pt(11)
    for s in doc.sections:
        s.left_margin = Inches(1.0); s.right_margin = Inches(1.0)
        s.top_margin = Inches(0.9); s.bottom_margin = Inches(0.9)
    add_page_number_footer(doc, footer_label)

    i, n = 0, len(lines)
    while i < n:
        line = lines[i]; stripped = line.strip()

        if not stripped:
            i += 1; continue

        if stripped == '[SALTO DE PÁGINA]':
            doc.add_page_break(); i += 1; continue

        if re.match(r'^-{3,}$', stripped):
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
            r = p.add_run('_'*60); r.font.color.rgb = LGREY
            i += 1; continue

        # Tabla
        if stripped.startswith('|') and i+1 < n and re.match(r'^\s*\|?[\s:|-]+\|?\s*$', lines[i+1]):
            header = [c.strip() for c in stripped.strip('|').split('|')]
            rows = []
            j = i+2
            while j < n and lines[j].strip().startswith('|'):
                rows.append([c.strip() for c in lines[j].strip().strip('|').split('|')]); j += 1
            empty_header = all(h == '' for h in header)
            if empty_header:
                # Bloque tipo "destinatario": sin bordes, 1ª columna en negrita
                table = doc.add_table(rows=0, cols=len(header))
                set_cell_no_borders(table)
                for row in rows:
                    cells = table.add_row().cells
                    for k in range(len(header)):
                        val = row[k] if k < len(row) else ''
                        cells[k].paragraphs[0].text = ''
                        add_runs(cells[k].paragraphs[0], val)
                        if k == 0:
                            for rr in cells[k].paragraphs[0].runs: rr.bold = True
            else:
                table = doc.add_table(rows=1, cols=len(header))
                table.style = 'Light Grid Accent 1'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
                hdr = table.rows[0].cells
                for k, h in enumerate(header):
                    hdr[k].paragraphs[0].text = ''
                    add_runs(hdr[k].paragraphs[0], h)
                    for rr in hdr[k].paragraphs[0].runs: rr.bold = True
                for row in rows:
                    cells = table.add_row().cells
                    for k in range(len(header)):
                        val = row[k] if k < len(row) else ''
                        cells[k].paragraphs[0].text = ''
                        add_runs(cells[k].paragraphs[0], val)
            doc.add_paragraph()
            i = j; continue

        # Títulos
        m = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if m:
            level = len(m.group(1)); txt = m.group(2).strip()
            p = doc.add_paragraph(); r = p.add_run(txt); r.bold = True
            if level == 1:
                r.font.size = Pt(19); r.font.color.rgb = BRAND
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(8)
            elif level == 2:
                r.font.size = Pt(14.5); r.font.color.rgb = BRAND
                p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
            elif level == 3:
                r.font.size = Pt(12); r.font.color.rgb = BRAND
                p.paragraph_format.space_before = Pt(7)
            else:
                r.font.size = Pt(11)
            i += 1; continue

        # Cita
        if stripped.startswith('>'):
            txt = stripped.lstrip('>').strip()
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
            add_runs(p, txt)
            for r in p.runs: r.italic = True; r.font.color.rgb = GREY
            i += 1; continue

        # Lista numerada
        m = re.match(r'^(\d+)\.\s+(.*)$', stripped)
        if m:
            p = doc.add_paragraph(style='List Number'); add_runs(p, m.group(2)); i += 1; continue

        # Viñetas
        if re.match(r'^[-*]\s+', stripped):
            p = doc.add_paragraph(style='List Bullet'); add_runs(p, re.sub(r'^[-*]\s+', '', stripped)); i += 1; continue

        # Párrafo normal (justificado)
        p = doc.add_paragraph(); add_runs(p, stripped)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        i += 1

    doc.save(docx_path); print('Generado:', docx_path)

DOCS = [
    ('docs/01_Comision_Calidad.md',     'docs/01_HemoPocket_Comision_Calidad.docx',     'HemoPocket · Comisión de Calidad'),
    ('docs/02_Jefatura_Servicio.md',    'docs/02_HemoPocket_Jefatura_Servicio.docx',    'HemoPocket · Servicio de Hematología'),
    ('docs/03_Direccion_Medica.md',     'docs/03_HemoPocket_Direccion_Medica.docx',     'HemoPocket · Dirección Médica'),
    ('docs/04_Informatica.md',          'docs/04_HemoPocket_Informatica.docx',          'HemoPocket · Sistemas de Información'),
    ('docs/05_Protocolo_Investigacion_Calidad.md', 'docs/05_HemoPocket_Protocolo_Investigacion.docx', 'HemoPocket · Protocolo de Investigación/Calidad'),
]

if __name__ == '__main__':
    import os
    for md, dx, label in DOCS:
        if os.path.exists(md):
            convert(md, dx, label)
    # Documentos base (proyecto y recomendaciones) si existen
    if os.path.exists('docs/PROYECTO_HemoPocket.md'):
        convert('docs/PROYECTO_HemoPocket.md', 'docs/PROYECTO_HemoPocket.docx', 'HemoPocket · Memoria del proyecto')
    if os.path.exists('docs/RECOMENDACIONES_aprobacion.md'):
        convert('docs/RECOMENDACIONES_aprobacion.md', 'docs/RECOMENDACIONES_aprobacion.docx', 'HemoPocket · Recomendaciones (uso interno)')
